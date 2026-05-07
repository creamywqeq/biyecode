from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

md_path = Path(r'c:\Doc\biyecode\thesis\毕业论文_完整合并版_基于3D技术的流场后处理模块开发.md')
template_path = Path(r'c:\Users\20677\Downloads\66eebb9f-53d4-45fe-9f71-d5e69371aa45.docx')
out_path = Path(r'c:\Doc\biyecode\thesis\毕业论文_直接填入学校模板版_基于3D技术的流场后处理模块开发.docx')

text = md_path.read_text(encoding='utf-8')
text = re.sub(r'```mermaid\n(.*?)\n```', lambda m: '\n[此处为 Mermaid 图，需在 Word 中替换为图片]\n', text, flags=re.S)
text = re.sub(r'```[a-zA-Z0-9_-]*\n(.*?)\n```', lambda m: '\n' + m.group(1).strip() + '\n', text, flags=re.S)

doc = Document(str(template_path))
sec = doc.sections[0]
# clear body content but preserve document-level styles/settings and section properties
body = doc._body._element
for child in list(body):
    if child.tag.endswith('}sectPr'):
        continue
    body.remove(child)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.5)

styles = doc.styles
for name in ['Normal', 'Body Text']:
    if name in styles:
        st = styles[name]
        st.font.name = '宋体'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        st.font.size = Pt(12)

def set_run_font(run, east='宋体', west='Times New Roman', size=12, bold=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    run.bold = bold

def add_para(content='', align=None, size=12, bold=False, first_indent=True, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = Pt(22)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_indent:
        fmt.first_line_indent = Pt(24)
    r = p.add_run(content)
    set_run_font(r, size=size, bold=bold)
    return p

def add_heading_school(content, level):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing = Pt(22)
    fmt.space_before = Pt(11)
    fmt.space_after = Pt(11)
    if level == 1:
        size = 18
    elif level == 2:
        size = 14
    else:
        size = 12
    r = p.add_run(content)
    set_run_font(r, east='黑体', size=size, bold=False)
    return p

def add_page_break():
    doc.add_page_break()

def is_table_block(lines, i):
    return i + 1 < len(lines) and lines[i].strip().startswith('|') and re.match(r'^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$', lines[i+1].strip())

def add_markdown_table(block):
    rows = []
    for idx, line in enumerate(block):
        if idx == 1:
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            set_run_font(run, size=10.5, bold=(r_idx == 0))

lines = text.splitlines()
i = 0
while i < len(lines):
    raw = lines[i]
    line = raw.strip()
    if not line:
        i += 1
        continue
    if 'page-break-after' in line:
        add_page_break()
        i += 1
        continue
    if is_table_block(lines, i):
        block = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i])
            i += 1
        add_markdown_table(block)
        continue
    if line.startswith('# '):
        add_heading_school(line[2:].strip(), 1)
    elif line.startswith('## '):
        add_heading_school(line[3:].strip(), 2)
    elif line.startswith('### '):
        add_heading_school(line[4:].strip(), 3)
    elif line.startswith('<p align="center">'):
        caption = re.sub(r'<[^>]+>', '', line).strip()
        add_para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=True, first_indent=False)
    else:
        line = re.sub(r'<sup>\[([^\]]+)\]</sup>', r'[\1]', line)
        line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        align = WD_ALIGN_PARAGRAPH.CENTER if line in ['Southwest University of Science and Technology', '本科毕业论文（设计）', '二〇二六年六月'] else WD_ALIGN_PARAGRAPH.LEFT
        add_para(line, align=align, size=12, bold=False, first_indent=(align != WD_ALIGN_PARAGRAPH.CENTER))
    i += 1

doc.save(str(out_path))
print(out_path)

